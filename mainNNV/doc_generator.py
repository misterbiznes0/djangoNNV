import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from django.conf import settings

def replace_placeholders(doc, context):
    """Замена плейсхолдеров в документе"""
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if key in paragraph.text:
                # Сохраняем стиль и форматирование
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(str(key), str(value))
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    if key in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(str(key), str(value))

def generate_attestation_educational(profile, request):
    """Генерация аттестационного листа (учебная практика)"""
    template_path = os.path.join(settings.BASE_DIR, 'templates', 'docs', 'attestation_educational.docx')
    
    if not os.path.exists(template_path):
        # Создаем шаблон, если его нет
        create_template_educational()
    
    doc = Document(template_path)
    
    # Подготовка контекста
    context = {
        '{{fio}}': profile.fio,
        '{{spec}}': profile.spec,
        '{{grupa}}': profile.grupa,
        '{{kurs}}': str(profile.kurs) if profile.kurs else '',
        '{{obuch}}': profile.obuch,
        '{{vid}}': profile.vid,
        '{{kod}}': profile.kod,
        '{{mesto}}': profile.mesto,
        '{{adress}}': profile.adress,
        '{{ruka}}': profile.ruka,
        '{{data}}': profile.date_begin.strftime('%d') if profile.date_begin else '',
        '{{data2}}': profile.date_begin.strftime('%m.%Y') if profile.date_begin else '',
        '{{data3}}': profile.date_finish.strftime('%d') if profile.date_finish else '',
        '{{data4}}': profile.date_finish.strftime('%m.%Y') if profile.date_finish else '',
        '{{god}}': profile.date_begin.strftime('%Y') if profile.date_begin else '',
        '{{god1}}': profile.date_finish.strftime('%Y') if profile.date_finish else '',
    }
    
    replace_placeholders(doc, context)
    
    # Сохраняем документ
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="attestation_educational_{profile.user.username}.docx"'
    doc.save(response)
    return response

def generate_attestation_industrial(profile, request):
    """Генерация аттестационного листа (производственная практика)"""
    template_path = os.path.join(settings.BASE_DIR, 'templates', 'docs', 'attestation_industrial.docx')
    
    if not os.path.exists(template_path):
        create_template_industrial()
    
    doc = Document(template_path)
    
    # Аналогичный контекст
    context = {
        '{{fio}}': profile.fio,
        '{{spec}}': profile.spec,
        '{{grupa}}': profile.grupa,
        '{{kurs}}': str(profile.kurs) if profile.kurs else '',
        '{{obuch}}': profile.obuch,
        '{{vid}}': profile.vid,
        '{{kod}}': profile.kod,
        '{{mesto}}': profile.mesto,
        '{{adress}}': profile.adress,
        '{{ruka}}': profile.ruka,
        '{{data}}': profile.date_begin.strftime('%d') if profile.date_begin else '',
        '{{data2}}': profile.date_begin.strftime('%m.%Y') if profile.date_begin else '',
        '{{data3}}': profile.date_finish.strftime('%d') if profile.date_finish else '',
        '{{data4}}': profile.date_finish.strftime('%m.%Y') if profile.date_finish else '',
        '{{god}}': profile.date_begin.strftime('%Y') if profile.date_begin else '',
        '{{god1}}': profile.date_finish.strftime('%Y') if profile.date_finish else '',
    }
    
    replace_placeholders(doc, context)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="attestation_industrial_{profile.user.username}.docx"'
    doc.save(response)
    return response

def generate_characteristic(profile, request):
    """Генерация характеристики"""
    template_path = os.path.join(settings.BASE_DIR, 'templates', 'docs', 'characteristic.docx')
    
    if not os.path.exists(template_path):
        create_characteristic_template()
    
    doc = Document(template_path)
    
    context = {
        '{{fio}}': profile.fio,
        '{{library}}': profile.library,
        '{{number}}': profile.number,
        '{{not_day_one}}': profile.not_day_one,
        '{{not_day_two}}': profile.not_day_two,
        '{{special}}': profile.special,
        '{{good}}': profile.good,
        '{{boss_organization}}': profile.boss_organization,
        '{{month_begin}}': profile.date_begin.strftime('%m') if profile.date_begin else '',
        '{{day_begin}}': profile.date_begin.strftime('%d') if profile.date_begin else '',
        '{{year_begin}}': profile.date_begin.strftime('%Y') if profile.date_begin else '',
        '{{month_finish}}': profile.date_finish.strftime('%m') if profile.date_finish else '',
        '{{day_finish}}': profile.date_finish.strftime('%d') if profile.date_finish else '',
        '{{year_finish}}': profile.date_finish.strftime('%Y') if profile.date_finish else '',
    }
    
    replace_placeholders(doc, context)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="characteristic_{profile.user.username}.docx"'
    doc.save(response)
    return response

def create_template_educational():
    """Создание шаблона аттестационного листа (учебная практика)"""
    # Эта функция должна создать шаблон на основе вашего файла
    # Сохраните ваш файл как attestation_educational.docx в templates/docs/
    pass

def create_template_industrial():
    """Создание шаблона производственной практики"""
    pass

def create_characteristic_template():
    """Создание шаблона характеристики"""
    pass