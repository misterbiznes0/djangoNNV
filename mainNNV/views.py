import os
import zipfile
from io import BytesIO

from docx import Document
import pymorphy3

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import login, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.http import HttpResponse
from django.shortcuts import render, redirect

from .models import Profile, DownloadHistory, GROUP_CHOICES, KURS_CHOICES, OBUCH_CHOICES, SPECIALTY_CHOICES


morph = pymorphy3.MorphAnalyzer()



def decline_fio(fio, case='gent'):
    if not fio:
        return ''

    parts = fio.strip().split()

    if len(parts) < 3:
        return fio

    tags = ['Surn', 'Name', 'Patr']
    result = []

    for word, tag in zip(parts[:3], tags):
        variants = [p for p in morph.parse(word) if tag in p.tag]
        parsed = variants[0] if variants else morph.parse(word)[0]
        declined = parsed.inflect({case})

        if declined:
            result.append(declined.word.capitalize())
        else:
            result.append(word)

    return ' '.join(result)


def mainNNV(request):
    return HttpResponse("app mainNNV")


def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)

        if form.is_valid():
            user = form.save()
            Profile.objects.get_or_create(user=user)
            login(request, user)
            messages.success(request, f'Добро пожаловать, {user.username}!')
            return redirect('/profile/')

        messages.error(request, "Ошибка регистрации. Проверьте введенные данные.")
        return render(request, "register.html", {"form": form})

    form = UserCreationForm()
    return render(request, "register.html", {"form": form})


def loginf(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)

        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)

            if user is not None:
                login(request, user)
                messages.success(request, f'Вы вошли как {username}')
                return redirect('/profile/')

        messages.error(request, 'Неправильное имя пользователя или пароль')

    form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})


def check_profile_errors(profile_obj):
    errors = []

    required_fields = {
        'fio': 'ФИО',
        'spec': 'Специальность',
        'grupa': 'Группа',
        'obuch': 'Форма обучения',
        'vid': 'Тип практики',
        'kod': 'Профессиональный модуль',
        'mesto': 'Наименование организации',
        'adress': 'Адрес организации',
        'library': 'Подразделение / место практики',
        'ruka': 'Руководитель практики',
        'boss_organization': 'Руководитель от организации',
        'date_begin': 'Дата начала практики',
        'date_finish': 'Дата окончания практики',
        'number': 'Количество выходов',
        'not_day_one': 'Пропущено дней',
        'not_day_two': 'По неуважительной причине',
        'good': 'Качество выполнения работы',
        'attestation_date': 'Дата аттестации',
    }

    for field, title in required_fields.items():
        value = getattr(profile_obj, field)
        if value in [None, '']:
            errors.append(title)

    if profile_obj.date_begin and profile_obj.date_finish:
        if profile_obj.date_begin > profile_obj.date_finish:
            errors.append('Дата начала практики не может быть позже даты окончания')

    return errors


@login_required
def profile(request):
    profile_obj, created = Profile.objects.get_or_create(user=request.user)

    if request.method == 'POST':
        profile_obj.fio = request.POST.get('fio', '')
        profile_obj.spec = request.POST.get('spec', '')
        profile_obj.grupa = request.POST.get('grupa', '')
        kurs_value = request.POST.get('kurs')
        profile_obj.kurs = int(kurs_value) if kurs_value else None
        profile_obj.obuch = request.POST.get('obuch', '')
        profile_obj.vid = request.POST.get('vid', '')
        profile_obj.kod = request.POST.get('kod', '')
        profile_obj.mesto = request.POST.get('mesto', '')
        profile_obj.adress = request.POST.get('adress', '')
        profile_obj.library = request.POST.get('library', '')
        profile_obj.ruka = request.POST.get('ruka', '')
        profile_obj.boss_organization = request.POST.get('boss_organization', '')
        profile_obj.date_begin = request.POST.get('date_begin') or None
        profile_obj.date_finish = request.POST.get('date_finish') or None
        profile_obj.number = request.POST.get('number') or ''
        profile_obj.not_day_one = request.POST.get('not_day_one') or ''
        profile_obj.not_day_two = request.POST.get('not_day_two') or ''
        # Поле special оставлено в модели для старых данных, но на сайте больше не вводится.
        # В новых шаблонах DOCX переменная {{special}} берётся из выбранной специальности spec.
        profile_obj.special = profile_obj.spec
        profile_obj.good = request.POST.get('good', '')
        profile_obj.attestation_date = request.POST.get('attestation_date') or None

        profile_obj.save()

        messages.success(request, "Данные успешно сохранены")
        return redirect('/profile/')

    profile_errors = check_profile_errors(profile_obj)
    history = DownloadHistory.objects.filter(user=request.user)[:10]

    return render(request, 'profile.html', {
        'profile': profile_obj,
        'profile_errors': profile_errors,
        'history': history,
        'group_choices': GROUP_CHOICES,
        'kurs_choices': KURS_CHOICES,
        'obuch_choices': OBUCH_CHOICES,
        'specialty_choices': SPECIALTY_CHOICES,
    })


def replace_placeholders(doc, context):
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        changed = False

        for key, value in context.items():
            if key in full_text:
                full_text = full_text.replace(key, str(value))
                changed = True

        if changed and full_text != paragraph.text:
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ''
            else:
                paragraph.text = full_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    changed = False

                    for key, value in context.items():
                        if key in full_text:
                            full_text = full_text.replace(key, str(value))
                            changed = True

                    if changed and full_text != paragraph.text:
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text
                            for run in paragraph.runs[1:]:
                                run.text = ''
                        else:
                            paragraph.text = full_text


def get_common_context(profile_obj):
    date_begin = profile_obj.date_begin
    date_finish = profile_obj.date_finish

    months = {
        1: 'января',
        2: 'февраля',
        3: 'марта',
        4: 'апреля',
        5: 'мая',
        6: 'июня',
        7: 'июля',
        8: 'августа',
        9: 'сентября',
        10: 'октября',
        11: 'ноября',
        12: 'декабря',
    }

    return {
        '{{fio}}': profile_obj.fio or '',
        '{{fio_rod}}': decline_fio(profile_obj.fio, 'gent'),
        '{{fio_dat}}': decline_fio(profile_obj.fio, 'datv'),
        '{{fio_vin}}': decline_fio(profile_obj.fio, 'accs'),

        '{{spec}}': profile_obj.spec or '',
        '{{grupa}}': profile_obj.grupa or '',
        '{{kurs}}': str(profile_obj.kurs) if profile_obj.kurs else '',
        '{{obuch}}': profile_obj.obuch or '',
        '{{vid}}': profile_obj.vid or '',
        '{{kod}}': profile_obj.kod or '',
        '{{mesto}}': profile_obj.mesto or '',
        '{{adress}}': profile_obj.adress or '',
        '{{library}}': profile_obj.library or '',
        '{{ruka}}': profile_obj.ruka or '',
        '{{boss_organization}}': profile_obj.boss_organization or '',
        '{{number}}': profile_obj.number or '',
        '{{not_day_one}}': profile_obj.not_day_one or '',
        '{{not_day_two}}': profile_obj.not_day_two or '',
        '{{special}}': profile_obj.spec or profile_obj.special or '',
        '{{good}}': profile_obj.good or '',

        '{{day_begin}}': str(date_begin.day) if date_begin else '',
        '{{month_begin}}': months.get(date_begin.month, '') if date_begin else '',
        '{{year_begin}}': str(date_begin.year) if date_begin else '',

        '{{day_finish}}': str(date_finish.day) if date_finish else '',
        '{{month_finish}}': months.get(date_finish.month, '') if date_finish else '',
        '{{year_finish}}': str(date_finish.year) if date_finish else '',

        '{{data}}': str(date_begin.day) if date_begin else '',
        '{{data2}}': months.get(date_begin.month, '') if date_begin else '',
        '{{data3}}': str(date_finish.day) if date_finish else '',
        '{{data4}}': months.get(date_finish.month, '') if date_finish else '',
        '{{god}}': str(date_begin.year) if date_begin else '',
        '{{god1}}': str(date_finish.year) if date_finish else '',
    }


def generate_document_bytes(profile_obj, template_filename):
    template_path = os.path.join(
        settings.BASE_DIR,
        'mainNNV',
        'templates',
        'docs',
        template_filename
    )

    if not os.path.exists(template_path):
        return None, template_path

    doc = Document(template_path)
    context = get_common_context(profile_obj)
    replace_placeholders(doc, context)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream, template_path


def download_single_document(request, template_filename, output_filename, history_name):
    profile_obj = Profile.objects.get(user=request.user)
    errors = check_profile_errors(profile_obj)

    if errors:
        messages.error(request, 'Заполните профиль до конца. Не заполнено: ' + ', '.join(errors))
        return redirect('/profile/')

    file_stream, template_path = generate_document_bytes(profile_obj, template_filename)

    if file_stream is None:
        return HttpResponse(f"Шаблон документа не найден по пути: {template_path}", status=404)

    DownloadHistory.objects.create(
        user=request.user,
        document_name=history_name,
        action='download'
    )

    response = HttpResponse(
        file_stream.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{output_filename}_{request.user.username}.docx"'

    return response


@login_required
def download_attestation_educational(request):
    return download_single_document(
        request,
        'attestation_educational.docx',
        'attestation_educational',
        'Аттестационный лист — учебная практика'
    )


@login_required
def download_attestation_industrial(request):
    return download_single_document(
        request,
        'attestation_industrial.docx',
        'attestation_industrial',
        'Аттестационный лист — производственная практика'
    )


@login_required
def download_characteristic(request):
    return download_single_document(
        request,
        'characteristic.docx',
        'characteristic',
        'Характеристика'
    )


@login_required
def download_all_documents_zip(request):
    profile_obj = Profile.objects.get(user=request.user)
    errors = check_profile_errors(profile_obj)

    if errors:
        messages.error(request, 'Заполните профиль до конца. Не заполнено: ' + ', '.join(errors))
        return redirect('/profile/')

    documents = [
        ('attestation_educational.docx', f'attestation_educational_{request.user.username}.docx'),
        ('attestation_industrial.docx', f'attestation_industrial_{request.user.username}.docx'),
        ('characteristic.docx', f'characteristic_{request.user.username}.docx'),
    ]

    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for template_filename, output_filename in documents:
            file_stream, template_path = generate_document_bytes(profile_obj, template_filename)

            if file_stream is None:
                return HttpResponse(f"Шаблон документа не найден по пути: {template_path}", status=404)

            zip_file.writestr(output_filename, file_stream.getvalue())

    zip_buffer.seek(0)

    DownloadHistory.objects.create(
        user=request.user,
        document_name='Все документы ZIP',
        action='download_zip'
    )

    response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="documents_{request.user.username}.zip"'

    return response