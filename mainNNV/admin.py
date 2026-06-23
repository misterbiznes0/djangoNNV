import os
import zipfile
from io import BytesIO
from datetime import datetime

from docx import Document

from django.contrib import admin, messages
from django.contrib.auth.admin import UserAdmin
from django.contrib.auth.models import User
from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.http import FileResponse, HttpResponse
from django.shortcuts import render, redirect
from django.urls import path
from django.utils.safestring import mark_safe

from .models import Profile, DownloadHistory, GROUP_CHOICES


def superuser_only(request):
    return (
        request.user.is_authenticated
        and request.user.is_active
        and request.user.is_superuser
    )


def protect_admin_view(request):
    if not superuser_only(request):
        raise PermissionDenied


admin.site.has_permission = superuser_only


class ProfileInline(admin.StackedInline):
    model = Profile
    can_delete = False
    verbose_name_plural = 'Профиль'
    extra = 0


class CustomUserAdmin(UserAdmin):
    inlines = [ProfileInline]

    list_display = (
        'username',
        'email',
        'is_superuser',
        'is_staff',
        'is_active',
        'get_profile_info',
    )

    list_editable = ('is_superuser', 'is_staff', 'is_active')

    def get_profile_info(self, obj):
        try:
            profile = obj.profile
            if profile.fio:
                return mark_safe('<span style="color: #28a745;">✓ Заполнен</span>')
            return mark_safe('<span style="color: #dc3545;">✗ Не заполнен</span>')
        except Profile.DoesNotExist:
            return mark_safe('<span style="color: #ffc107;">⚠ Нет профиля</span>')

    get_profile_info.short_description = 'Профиль'


admin.site.unregister(User)
admin.site.register(User, CustomUserAdmin)


@admin.register(Profile)
class ProfileAdmin(admin.ModelAdmin):
    list_display = ('user', 'fio', 'spec', 'grupa', 'kurs', 'vid')
    list_filter = ('grupa', 'vid', 'kurs', 'obuch')
    search_fields = ('fio', 'spec', 'grupa', 'user__username')


@admin.register(DownloadHistory)
class DownloadHistoryAdmin(admin.ModelAdmin):
    list_display = ('user', 'document_name', 'action', 'created_at')
    list_filter = ('action', 'created_at')
    search_fields = ('user__username', 'document_name')
    readonly_fields = ('user', 'document_name', 'action', 'created_at')


def replace_placeholders(doc, context):
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        changed = False

        for key, value in context.items():
            if key in full_text:
                full_text = full_text.replace(key, str(value))
                changed = True

        if changed:
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ''
            else:
                paragraph.text = full_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    changed = False

                    for key, value in context.items():
                        if key in full_text:
                            full_text = full_text.replace(key, str(value))
                            changed = True

                    if changed:
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text
                            for run in paragraph.runs[1:]:
                                run.text = ''
                        else:
                            paragraph.text = full_text


def get_profile_errors(profile):
    errors = []

    required_fields = {
        'fio': 'ФИО',
        'spec': 'Специальность',
        'grupa': 'Группа',
        'kurs': 'Курс',
        'obuch': 'Форма обучения',
        'vid': 'Тип практики',
        'kod': 'Профессиональный модуль',
        'mesto': 'Организация',
        'adress': 'Адрес',
        'library': 'Подразделение',
        'ruka': 'Руководитель практики',
        'boss_organization': 'Руководитель от организации',
        'date_begin': 'Дата начала',
        'date_finish': 'Дата окончания',
        'number': 'Количество выходов',
        'not_day_one': 'Пропущено дней',
        'not_day_two': 'По неуважительной причине',
        'special': 'Специальность на практике',
        'good': 'Качество выполнения работы',
        'attestation_date': 'Дата аттестации',
    }

    for field, title in required_fields.items():
        value = getattr(profile, field)
        if value in [None, '']:
            errors.append(title)

    if profile.date_begin and profile.date_finish and profile.date_begin > profile.date_finish:
        errors.append('Дата начала позже даты окончания')

    return errors


    def get_context(profile):
        date_begin = profile.date_begin
        date_finish = profile.date_finish

        MONTHS = {
            1: 'января',
            2: 'февраля',
            3: 'марта',
            4: 'апреля',
            5: 'мая',
            6: 'июня',
            7: 'июля',
            8: 'августа',
            9: 'сентября',
            10: 'октября',
            11: 'ноября',
            12: 'декабря',
        }

        return {
            '{{fio}}': profile.fio or '',
            '{{spec}}': profile.spec or '',
            '{{grupa}}': profile.grupa or '',
            '{{kurs}}': str(profile.kurs) if profile.kurs else '',
            '{{obuch}}': profile.obuch or '',
            '{{vid}}': profile.vid or '',
            '{{kod}}': profile.kod or '',
            '{{mesto}}': profile.mesto or '',
            '{{adress}}': profile.adress or '',
            '{{ruka}}': profile.ruka or '',
            '{{library}}': profile.library or '',
            '{{number}}': profile.number or '',
            '{{not_day_one}}': profile.not_day_one or '',
            '{{not_day_two}}': profile.not_day_two or '',
            '{{special}}': profile.special or '',
            '{{good}}': profile.good or '',
            '{{boss_organization}}': profile.boss_organization or '',

            # 🔥 НОВЫЕ КРАСИВЫЕ ДАТЫ
            '{{day_begin}}': date_begin.strftime('%d') if date_begin else '',
            '{{month_begin}}': MONTHS.get(date_begin.month) if date_begin else '',
            '{{year_begin}}': date_begin.strftime('%Y') if date_begin else '',

            '{{day_finish}}': date_finish.strftime('%d') if date_finish else '',
            '{{month_finish}}': MONTHS.get(date_finish.month) if date_finish else '',
            '{{year_finish}}': date_finish.strftime('%Y') if date_finish else '',
        }
def get_context(profile):
    date_begin = profile.date_begin
    date_finish = profile.date_finish

    months = {
        1: 'января',
        2: 'февраля',
        3: 'марта',
        4: 'апреля',
        5: 'мая',
        6: 'июня',
        7: 'июля',
        8: 'августа',
        9: 'сентября',
        10: 'октября',
        11: 'ноября',
        12: 'декабря',
    }

    return {
        '{{fio}}': profile.fio or '',
        '{{spec}}': profile.spec or '',
        '{{grupa}}': profile.grupa or '',
        '{{kurs}}': str(profile.kurs) if profile.kurs else '',
        '{{obuch}}': profile.obuch or '',
        '{{vid}}': profile.vid or '',
        '{{kod}}': profile.kod or '',
        '{{mesto}}': profile.mesto or '',
        '{{adress}}': profile.adress or '',
        '{{library}}': profile.library or '',
        '{{ruka}}': profile.ruka or '',
        '{{boss_organization}}': profile.boss_organization or '',
        '{{number}}': profile.number or '',
        '{{not_day_one}}': profile.not_day_one or '',
        '{{not_day_two}}': profile.not_day_two or '',
        '{{special}}': profile.special or '',
        '{{good}}': profile.good or '',

        '{{day_begin}}': str(date_begin.day) if date_begin else '',
        '{{month_begin}}': months.get(date_begin.month, '') if date_begin else '',
        '{{year_begin}}': str(date_begin.year) if date_begin else '',

        '{{day_finish}}': str(date_finish.day) if date_finish else '',
        '{{month_finish}}': months.get(date_finish.month, '') if date_finish else '',
        '{{year_finish}}': str(date_finish.year) if date_finish else '',

        '{{data}}': str(date_begin.day) if date_begin else '',
        '{{data2}}': months.get(date_begin.month, '') if date_begin else '',
        '{{data3}}': str(date_finish.day) if date_finish else '',
        '{{data4}}': months.get(date_finish.month, '') if date_finish else '',
        '{{god}}': str(date_begin.year) if date_begin else '',
        '{{god1}}': str(date_finish.year) if date_finish else '',
    }
def generate_docx_bytes(profile, template_filename):
    template_path = os.path.join(
        settings.BASE_DIR,
        'mainNNV',
        'templates',
        'docs',
        template_filename
    )

    if not os.path.exists(template_path):
        return None, template_path

    doc = Document(template_path)
    replace_placeholders(doc, get_context(profile))

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    return stream.getvalue(), template_path


def custom_admin_index(request):
    protect_admin_view(request)

    user_count = User.objects.count()
    active_users = User.objects.filter(is_active=True).count()
    superusers = User.objects.filter(is_superuser=True).count()
    recent_users = User.objects.order_by('-date_joined')[:5]

    templates_dir = os.path.join(settings.BASE_DIR, 'mainNNV', 'templates', 'docs')
    templates = []
    template_count = 0

    if os.path.exists(templates_dir):
        for file in os.listdir(templates_dir):
            if file.endswith('.docx'):
                template_count += 1
                file_path = os.path.join(templates_dir, file)
                stat = os.stat(file_path)
                templates.append({
                    'name': file,
                    'size': stat.st_size,
                    'modified': datetime.fromtimestamp(stat.st_mtime),
                })

    return render(request, 'admin/custom_admin.html', {
        'user_count': user_count,
        'active_users': active_users,
        'superusers': superusers,
        'recent_users': recent_users,
        'templates': templates,
        'template_count': template_count,
        'title': 'Панель управления',
    })


def template_list(request):
    protect_admin_view(request)

    templates_dir = os.path.join(settings.BASE_DIR, 'mainNNV', 'templates', 'docs')
    templates = []

    if os.path.exists(templates_dir):
        for file in os.listdir(templates_dir):
            if file.endswith('.docx'):
                file_path = os.path.join(templates_dir, file)
                stat = os.stat(file_path)
                templates.append({
                    'name': file,
                    'size': stat.st_size,
                    'modified': datetime.fromtimestamp(stat.st_mtime),
                })

    return render(request, 'admin/template_list.html', {
        'templates': templates,
        'title': 'Шаблоны документов',
    })


def upload_template(request):
    protect_admin_view(request)

    if request.method == 'POST' and request.FILES.get('template_file'):
        template_file = request.FILES['template_file']

        if template_file.name.endswith('.docx'):
            templates_dir = os.path.join(settings.BASE_DIR, 'mainNNV', 'templates', 'docs')
            os.makedirs(templates_dir, exist_ok=True)

            file_path = os.path.join(templates_dir, template_file.name)

            with open(file_path, 'wb+') as f:
                for chunk in template_file.chunks():
                    f.write(chunk)

            messages.success(request, f'Шаблон "{template_file.name}" загружен')
        else:
            messages.error(request, 'Можно загружать только DOCX')

    return redirect('admin:template_list')


def delete_template(request, template_name):
    protect_admin_view(request)

    path_file = os.path.join(settings.BASE_DIR, 'mainNNV', 'templates', 'docs', template_name)

    if os.path.exists(path_file):
        os.remove(path_file)
        messages.success(request, f'Удалено: {template_name}')

    return redirect('admin:template_list')


def download_template(request, template_name):
    protect_admin_view(request)

    path_file = os.path.join(settings.BASE_DIR, 'mainNNV', 'templates', 'docs', template_name)

    if os.path.exists(path_file):
        response = FileResponse(open(path_file, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{template_name}"'
        return response

    messages.error(request, 'Файл не найден')
    return redirect('admin:template_list')


def generate_documents_page(request):
    protect_admin_view(request)

    return render(request, 'admin/generate_documents.html', {
        'groups': GROUP_CHOICES,
        'title': 'Генерация документов',
    })


def generate_documents_zip(request):
    protect_admin_view(request)

    selected_group = request.POST.get('group', 'all')

    if selected_group == 'all':
        profiles = Profile.objects.select_related('user').all().order_by('grupa', 'fio')
        zip_name = 'college_documents.zip'
    else:
        profiles = Profile.objects.select_related('user').filter(grupa=selected_group).order_by('fio')
        zip_name = f'{selected_group}_documents.zip'

    if not profiles.exists():
        messages.error(request, 'Нет профилей для выбранной группы')
        return redirect('admin:generate_documents_page')

    documents = [
        ('attestation_educational.docx', 'attestation_educational.docx'),
        ('attestation_industrial.docx', 'attestation_industrial.docx'),
        ('characteristic.docx', 'characteristic.docx'),
    ]

    errors = []
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for profile in profiles:
            profile_errors = get_profile_errors(profile)

            student_name = profile.fio or profile.user.username
            group_name = profile.grupa or 'Без группы'

            if profile_errors:
                errors.append(f'{group_name} / {student_name}: не заполнено — {", ".join(profile_errors)}')
                continue

            safe_student_name = student_name.replace('/', '-').replace('\\', '-')
            safe_group_name = group_name.replace('/', '-').replace('\\', '-')

            for template_filename, output_filename in documents:
                doc_bytes, template_path = generate_docx_bytes(profile, template_filename)

                if doc_bytes is None:
                    errors.append(f'Не найден шаблон: {template_path}')
                    continue

                zip_file.writestr(
                    f'{safe_group_name}/{safe_student_name}/{output_filename}',
                    doc_bytes
                )

            DownloadHistory.objects.create(
                user=profile.user,
                document_name='Массовая генерация документов',
                action='admin_group_zip'
            )

        if errors:
            zip_file.writestr('errors.txt', '\n'.join(errors))

    zip_buffer.seek(0)

    response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{zip_name}"'

    return response


original_urls = admin.site.get_urls


def new_urls():
    return [
        path('', admin.site.admin_view(custom_admin_index), name='index'),
        path('templates/', admin.site.admin_view(template_list), name='template_list'),
        path('templates/upload/', admin.site.admin_view(upload_template), name='upload_template'),
        path('templates/delete/<str:template_name>/', admin.site.admin_view(delete_template), name='delete_template'),
        path('templates/download/<str:template_name>/', admin.site.admin_view(download_template), name='download_template'),

        path('generate-documents/', admin.site.admin_view(generate_documents_page), name='generate_documents_page'),
        path('generate-documents/download/', admin.site.admin_view(generate_documents_zip), name='generate_documents_zip'),
    ] + original_urls()


admin.site.get_urls = new_urls
admin.site.index_template = 'admin/custom_admin.html'

admin.site.site_header = 'Панель управления'
admin.site.site_title = 'Админка'
admin.site.index_title = 'Добро пожаловать'